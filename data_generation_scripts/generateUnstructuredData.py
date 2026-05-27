"""
Generates unstructured companion data files for InsightForge.
Files contain a mix of:
  - Data that matches the Superstore CSV (regions, products, customers)
  - Data that does NOT match (fictional markets, different companies)
  - Noise: typos, OCR artifacts, inconsistent formatting, placeholders
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib import colors
from reportlab.lib.units import inch

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data")
os.makedirs(OUT, exist_ok=True)


# ─────────────────────────────────────────────
#  TXT FILES
# ─────────────────────────────────────────────

def write_txt(filename, content):
    path = os.path.join(OUT, filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"  Created: {filename}")


# TXT 1 — Quarterly Sales Report (matches Superstore regions/categories + noise)
write_txt("quarterly_sales_report_Q4_2017.txt", """
QUARTERLY SALES REPORT - Q4 2017
TechRetail Corporation -- Internal Use Only
Prepared by: Regional Analytics Team
Date: Januery 3, 2018  [NOTE: please verify date before sending]

=====================================================================
EXECUTIVE SUMMARY
=====================================================================
Total Q4 revenue reached $  847,932.00 across all regons.
West Region perfomed best with sales of $312,445 - up 14% YoY.
The South region showed declining trend w/ total of $198,221.
Centra region: $176,833 (flat growth vs Q3).
East Region: $160,433 - BELOW target by ~8%.

NOTE: some figures unverified -- pending reconciliation from finance

=====================================================================
REGIONAL BREAKDOWN
=====================================================================

WEST REGION:
  Technology sales : $145,332
  Furniture        : $98,211
  Office Supplies  : $68,902
  Top state        : California ($189,443)
  2nd              : Washington ($72,002)
  3rd              : Oregon ($51,000)  [figure approximate]

EAST REGION
  Technology       : $67,443
  Furniture        : $52,100
  Office supplies  : $40890            <- missing $ sign in source
  Key states       : New York, Pennsylvania, Ohio
  Nw York alone    : $88,332

SOUTH REGIO:                           <- truncated in source system
  Tech             : $89,002
  Furn             : $65,441
  OfficeSup        : $43,778
  Note: Texas data may be incomplete - systems outage Dec 12-14

CENTRAL REGION:
  Technology       : $72,113
  Furniture        : $60,002
  Office supplies  : $44,718
  Illinois contrib : $65,331

=====================================================================
PRODUCT PERFORMANCE
=====================================================================
Top Sub-Categories by Revenue (Q4 2017):
 1.  Phones          $98,331    [verify against CRM]
 2.  Chairs          $87,003
 3.  Storage         $61,002
 4.  Tables          $58,441    <- includes clearance / discount items
 5.  Binders         $43,112
 6.  Machines        $41,009
 7.  Accessories     $38,771
 8.  Copiers         $35,002
 9.  Bookcases       $31,441
10.  Appliences      $28,002    <- spelling error in source system, do not correct

=====================================================================
CUSTOMER SEGMENTS
=====================================================================
Consumer     :  46.2%  ($391,943)
Corporate    :  32.7%  ($277,224)
Home Office  :  21.1%  ($178,765)

Note: percentages dont add to 100 due to rounding -- see finance

=====================================================================
SHIPPING PERFORMANCE
=====================================================================
Standard Class :  58% of orders   avg 5.2 days
Second Class   :  20% of orders   avg 3.1 days
First Class    :  15% of orders   avg 2.0 days
Same Day       :   7% of orders   avg 0.9 dyas   <- typo in log

ISSUES FLAGGED:
  - 47 orders delayed in South region (holiday backlog)
  - Same-day cost overrun: $12,441 vs budget $9,000
  - 3 escalations pending: CG-12520, AA-10480, BH-11710

=====================================================================
CONCERNS & RECOMMENDATIONS
=====================================================================
1. Discount abuse in Furniture -- avg 42% vs target 20%
   Action: Review with sales leads by Jan 15
2. Negative profit on Tables (-$8,933) needs investigation
3. APAC expansion data NOT in this report (see separate deck)
4. Q1 2018 targets TBD -- budget meeting Feb 2

TODO: Get VP Sales sign-off before wider distribution
[PLACEHOLDER -- insert CFO comment here before sending]
[PLACEHOLDER -- attach appendix A: raw regional data]

---END OF REPORT---
Distribution: Internal Only | Classification: CONFIDENTIAL
""")


# TXT 2 — Customer Feedback Notes (partial match + misspellings + gaps)
write_txt("customer_feedback_notes.txt", """
CUSTOMER FEEDBACK COMPILATION
Source: CRM export + manual transcription from call recordings
Period: October -- December 2017
Transcribed by: Support Team  (raw -- errors may exist)
=========================================================

ID: FB-001
Customer  : Clair Gute     [possible correct spelling: Claire Gute]
Segment   : Consumer
Region    : South
Date      : 10/14/2017
Feedback  : "recieved my bookcase damaged. corners dented, one shelf cracked.
             took 3 weeks with Second Clas shipping. want refund or replacement!!"
Sentiment : NEGATIVE
Product   : Bush Somerset Collection Bookcas   <- truncated
Follow-up : Replacement shipped 10/18. Customer satisfied.
---------------------------------------------------------

ID: FB-002
Custmer   : Darrin Van Huff          <- typo in transcription
Region    : West
Date      : Nov 3 '17
Feedback  : Great service!! ordered Hon chairs for entire office. arrived on time
            good quality price was resonable. Will order again definitely
Segment   : Coporate                 <- typo
Sentiment : POSITIV                  <- incomplete word
Product   : Hon Deluxe Fabric Upholstered Stacking Chairs
---------------------------------------------------------

ID: FB-003
Customer  : Sean O'Donnell
Region    : Suth                     <- OCR error, should be South
Date      : 11/17/2017
Feedback  : "storage box completely wrong colour. ordered black got grey.
             qty ordered 2 but received 1. VERY DISAPPOINTING. 3rd issue this yr"
Issues    : Wrong item, wrong qty
Sentiment : Negative
Return    : YES   ref# RTN-2017-0443
---------------------------------------------------------

ID: FB-004
Custoer   : Maris Freeman            <- garbled OCR
Region    : East
Dat       : 12 / 01 / 2017           <- inconsistent date format
Feedback  : Phone call -- customer very happy with Logitech MK520 keyboard.
            Best keyboard she has used. Recommended to 4 colleagues. No issues.
Sentiment : Positive
NPS Score : 9/10
---------------------------------------------------------

ID: FB-005
Customer  : AARON HAWKINS
Segment   : Home Office
Region    : Centra                   <- OCR truncation
Date      : 2017-12-22
Feedback  : staples and paper arrived quickly but stapler itself was missing.
            checked invoice -- was charged for it. pls credit or resend
Missing   : Stapler
Invoice#  : CA-2017-8819
Status    : OPEN -- pending warehouse investigation
---------------------------------------------------------

ID: FB-006
Customer  : Somebody McTest          <- TEST RECORD -- DELETE BEFORE PRODUCTION
Region    : N/A
Date      : 00/00/0000
Feedback  : Lorem ipsum dolor sit amet, consectetur adipiscing elit.
Sentiment : unknown
Status    : TEST DATA -- IGNORE
---------------------------------------------------------

ID: FB-007
Customer  : Heinrich Braun           <- NOTE: European customer, different system
Region    : EMEA - Germany           <- NOT in Superstore dataset
Date      : 12/05/2017
Feedback  : Sehr gut! Schnelle Lieferung, Qualitat top.
            [Translation needed -- auto-translate failed]
Sentiment : Positive (inferred)
Status    : Escalate to EU team for follow-up
---------------------------------------------------------

ID: FB-008
Customer  : Priya Nair
Region    : West
Segment   : Corporate
Date      : December 19, 2017        <- different date format again
Feedback  : Bulk order of 200 binders arrived short by 23 units. Invoice correct
            but physical count wrong. Warehouse error suspected.
Invoice#  : CA-2017-9942
Status    : Pending credit note
NPS Score : 6
---------------------------------------------------------

ADDITIONAL NOTES (unverified, raw):
  - Several customers complained abot holiday delays in South & Central
  - Furniture returns up ~12% Q4 vs Q3 -- possible supplier quality issue
  - Technology items lowest complaint rate (2.1%)
  - One VIP Corporate customer West placed $45,000 order Dec 15 -- no issues
  - EU feedback NOT in this file -- separate system
  - TODO: NPS analysis by region and product category
  - TODO: clean up TEST records before analysis

[INCOMPLETE -- 340 remaining feedback records in CRM system]
[LAST UPDATED: ???  -- timestamp missing from export]
""")


# TXT 3 — Market Analysis (external + internal mix, junk sections, noise)
write_txt("market_analysis_2018.txt", """
MARKET ANALYSIS REPORT -- US RETAIL SECTOR
InsightForge Business Intelligence Unit
FY2018 Outlook
DRAFT v0.3  <<< NOT FOR EXTERNAL DISTRIBUTION >>>

=========================================================
EXECUTIVE BRIEF
=========================================================
The US retail market continues too evolve rapidly. Omnichannel strategies
dominate and companies that fail to adapt risk obsoletion. This report
analyzses key trends, competitive landscape, and strategic opportunityes.

Note: Data sourced from multiple vendors -- consistency NOT guaranteed
      Some projections use 2016 base -- update with 2017 actuals when avail.

=========================================================
SECTION 1: MARKET SIZE & GROWTH
=========================================================
Total US Office Products Market (2017): $14.2 billion
YoY Growth  : 3.2%  (vs 2.8% in 2016)
2018 Proj.  : $14.8 - $15.1 billion  (range reflects uncertainty)

Segment Breakdown:
  Office Supplies  : 38% market share   [Source: IBISWorld, possibly outdated]
  Furniture        : 29% market share
  Technology       : 33% market share
  [Note: sums to 100 but cross-check with Nielsen data -- discrepancy noted]

US Regional Share:
  West Coast  : 28%
  Northeast   : 25%
  Southeast   : 22%
  Midwest     : 15%
  Southwest   : 10%
  [figures may not sum -- rounding issue]

=========================================================
SECTION 2: COMPETITIVE LANDSCAPE
=========================================================
1. Staples Inc.
   Revenue 2017 : ~$10.2B  (estimate, not confirmed)
   Strength     : Brand, B2B contracts
   Weakness     : High cost, slow digital transformation

2. Office Depot / OfficeMax
   Revenue      : ~$11.0B  (includes CompuCom)
   Note: Merger integration ongoing -- operational disruption expected Q1 2018

3. Amazon Business  <<< MAJOR THREAT >>>
   Est 2017 GMV : $10B+
   Differentiator: Price, delivery speed, selection breadth
   Risk to us   : HIGH -- price matching not sustainable long-term

4. W.W. Grainger
   Focus: Industrial/MRO -- partial overlap in supplies

5. TechRetail Corp (our position):
   Revenue 2017 : ~$2.3M   [from internal report -- verify with finance]
   Strengths    : Customer service, product quality, relationships
   Gaps         : Digital presence, pricing vs Amazon, analytics maturity

=========================================================
SECTION 3: CONSUMER TRENDS (2017-2018)
=========================================================
a) REMOTE WORK ACCELERATION
   Home Office segment YoY growth: +18%
   Ergonomic furniture demand: significant increase (exact % TBD)
   Tech accessories top growth: webcams, headsets, monitors

b) SUSTAINABILITY DEMAND
   Green product enquiries   : +34%
   Eco-cert paper requests   : increasing
   Packaging complaints      : rising (our weak area)

c) PRICE SENSITIVITY
   Expected discount range   : 15-25%
   Online price comparison   : >80% of B2C customers
   Brand loyalty declining in commodity items (pens, paper, toner)

d) DELIVERY EXPECTATIONS
   2-day delivery expectation : 70% of surveyed customers
   Same-day demand growing but avg cost $18/order -- unprofitable at scale
   [Our same-day costs over budget in Q4 -- see ops report]

=========================================================
SECTION 4: OUR PERFORMANCE vs MARKET
=========================================================
TechRetail 2017 Summary:
  Revenue                : $2,297,200
  Gross Margin Overall   : ~12.5%
  Technology margin      : 17%   (best)
  Furniture margin       : 6.2%  (worst)
  Customer retention     : 68%   (industry avg 72%) -- BELOW average

Problem Areas:
  1. Tables sub-category     : -4.2% margin (loss-making)
  2. South region            : below national avg satisfaction
  3. Furniture avg discount  : 22% vs industry norm 15%
  4. Mobile website          : 34% bounce rate (industry benchmark 22%)

=========================================================
SECTION 5: STRATEGIC RECOMMENDATIONS   [DRAFT]
=========================================================
SHORT TERM Q1-Q2 2018:
  - Dynamic pricing rollout for Technology
  - Supplier renegotiation for Furniture (Tables specifically)
  - Targeted marketing push: Central Region (untapped upside)
  - Fix mobile feedback form bug (reported by support Q4)
  - Audit discount approval process

MEDIUM TERM H2 2018:
  - Same-day delivery courier partnership (pilot West region first)
  - Corporate loyalty programme launch
  - Data analytics platform upgrade (current tools insufficient)
  - EMEA market feasibility study  [NOT current scope -- separate workstream]

LONG TERM 2019+:
  - AI-powered demand forecasting
  - Warehouse automation (West Region priority)
  - Potential South regional competitor acquisition (TBD)

=========================================================
SECTION 6: RISK REGISTER
=========================================================
HIGH   : Amazon price competition -- unsustainable to match
MEDIUM : Supplier concentration (top 3 = 67% of COGS)
MEDIUM : Analytics/tech talent retention
MEDIUM : Tariff impact on imported technology products (2018 uncertainty)
LOW    : GDPR (minimal US ops impact unless EU expansion proceeds)
UNKNOWN: Macro recession risk -- leading indicators mixed

=========================================================
APPENDICES
=========================================================
Appendix A: Raw survey data              [ATTACHED SEPARATELY]
Appendix B: Competitor pricing analysis  [REDACTED -- legal hold]
Appendix C: Financial projections        [SEE CFO DECK -- NOT THIS FILE]
Appendix D: EMEA market data             [OUT OF SCOPE -- REMOVED]

DOCUMENT STATUS : DRAFT
Last edited     : [DATE NOT UPDATED -- TEMPLATE PLACEHOLDER]
Author          : [AUTHOR PLACEHOLDER]
Reviewer        : TBD
Approved        : NOT APPROVED
""")


# TXT 4 — Internal ops memo with heavy noise and partial data
write_txt("operations_memo_Dec2017.txt", """
INTERNAL MEMO
TO   : All Regional Ops Managers
FROM : Head of Operations
DATE : Decmber 28, 2O17   <- note: O vs 0 OCR issue
RE   : End-of-Year Ops Review & January Priorities

---

This memo covers outstanding issues from Q4 and sets priorities for Jan 2018.
Plese acknowledge receipt by repling to this email.

OUTSTANDING ISSUES:
===================

1. INVENTORY DISCREPANCIES
   West warehouse: 142 units Staple unaccounted (item# OFF-ST-10000760)
   South warehouse: Tables overstock -- 38 units beyond forecast (margin -ve)
   Central: OK
   East: Minor variance $1,200 -- within tolerance

2. RETURNS PROCESSING BACKLOG
   Current backlog: 89 returns unprocessed (vs SLA of 48hr)
   Main categories: Furniture (51), Technology (23), Office Supplies (15)
   Root cause: 2 staff on leave, temp cover not arranged
   ACTION: Regional managers to provide 1 person each until cleared

3. SHIPPING CARRIER ISSUE
   FedEx contract expires Jan 31 2018
   Renewal negotiations IN PROGRESS -- do not discuss externally
   Backup carrier identified: UPS (rates 8% higher -- seeking approval)
   Same-day carrier for NYC area: Deliv (pilot -- results mixed so far)

4. SYSTEM OUTAGE LOG (Q4):
   Dec 12 14:00-18:30 : South region order system DOWN (472 orders delayed)
   Dec 23 09:00-09:45 : Payment gateway intermittent (12 failed transactions)
   Nov 08 02:00-03:00 : Scheduled maintenance (all regions -- no impact)

5. CUSTOMER ESCALATIONS (unresolved):
   ESC-2017-0891 : Large furniture order dispute -- Corporate client -- West
   ESC-2017-0934 : Repeated wrong item delivery -- Consumer -- South
   ESC-2017-0971 : Billing error $3,400 overcharge -- Corporate -- East
   [See CRM for full details -- do not discuss by email]

UPCOMING JANUARY PRIORITIES:
=============================
Week 1 (Jan 1-5)  : Year-end stock count (ALL warehouses mandatory)
Week 2 (Jan 8-12) : Carrier contract renewal decision needed
Week 3 (Jan 15)   : Discount policy review meeting (Sales + Ops)
Week 4 (Jan 22)   : Q1 demand forecast submission to procurement

HR NOTE: 3 new hires starting Jan 8 (2 warehouse, 1 analytics).
         Onboarding materials in Sharepoint -- link: [BROKEN LINK PLACEHOLDER]

BUDGET REMINDER:
Q1 ops budget NOT yet approved. Work within Q4 actuals until notified.
Any spend >$5,000 requires SVP approval until budget confirmed.

---
CONFIDENTIAL -- INTERNAL USE ONLY
Do not forward outside TechRetail Corp.
If received in error please delete and notify sender.

[AUTO-SIGNATURE FAILED -- CONTACT IT]
""")

print("  TXT files done.")


# ─────────────────────────────────────────────
#  DOCX FILES
# ─────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)
    return p


# DOCX 1 — Annual Report 2017
doc1 = Document()
doc1.add_heading("TechRetail Corporation", 0)
doc1.add_heading("ANNUAL REPORT 2017", 1)
add_para(doc1, "For Internal Distribution Only | Confidential", color=(180, 0, 0))
doc1.add_paragraph()

doc1.add_heading("MESSAGE FROM THE CEO", 1)
doc1.add_paragraph(
    "Dear Team,\n\n"
    "2017 was a year of mixed results for TechRetail Corporation. We achieved revenue growth "
    "of 6.2% year-on-year, reaching $2,297,200 in total sales across our four operating regions. "
    "However, profitability remained under pressure, particularly in the Furniture category, "
    "where aggressive discounting eroded margins.\n\n"
    "Our Technology segment was the standout performer, delivering 17% gross margin and strong "
    "demand across all regions. The West Region continued its dominant position, contributing "
    "approximately 37% of total revenue.\n\n"
    "Challenges remain. Customer retention at 68% is below the industry benchmark of 72%. "
    "We lost ground in the South Region and must address operational inefficiencies that led "
    "to Q4 shipping delays. The competitive threat from online platforms intensifies.\n\n"
    "Looking to 2018, our priorities are: margin recovery in Furniture, digital capability "
    "investment, and improved customer experience. I thank every member of our team for their "
    "dedication this year.\n\n"
    "J. Matthews\nCEO, TechRetail Corporation\n[SIGNATURE PLACEHOLDER]"
)

doc1.add_heading("FINANCIAL HIGHLIGHTS 2017", 1)
table = doc1.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Metric"
hdr[1].text = "2017"
hdr[2].text = "2016"
rows = [
    ("Total Revenue", "$2,297,200", "$2,162,300"),
    ("Gross Profit", "$286,397", "$280,100"),
    ("Gross Margin %", "12.5%", "13.0%"),
    ("Technology Revenue", "$836,154", "$765,200"),
    ("Furniture Revenue", "$741,999", "$698,100"),
    ("Office Supplies Revenue", "$719,047", "$699,000"),
    ("West Region Revenue", "$725,457", "$680,000"),
    ("East Region Revenue", "$678,781", "$635,000"),
    ("South Region Revenue", "$501,239", "$498,000 [NOTE: verify]"),
    ("Central Region Revenue", "$391,723", "$349,300"),
    ("Total Orders", "9,994", "9,426"),
    ("Avg Order Value", "$229.85", "$229.38"),
]
for r in rows:
    row = table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

doc1.add_paragraph()
add_para(doc1, "NOTE: 2016 figures are unaudited estimates. Final audit report pending.", color=(200, 100, 0))

doc1.add_heading("REGIONAL PERFORMANCE", 1)
doc1.add_paragraph(
    "West Region performed strongly across all three product categories. California remained "
    "the single largest state by revenue. Investment in same-day delivery capability in the "
    "West contributed to customer satisfaction improvements, albeit at higher cost.\n\n"
    "East Region grew modestly. New York City remained the top metropolitan market. Corporate "
    "segment engagement increased following a targeted outreach campaign in Q3.\n\n"
    "South Region was a concern. Mississippi and Arkansas underperformed significantly against "
    "targets. A December systems outage (Dec 12-14) disrupted approximately 472 orders, "
    "damaging customer satisfaction scores in the region for the quarter.\n\n"
    "Central Region demonstrated the highest growth rate (12.2% YoY) from a lower base. "
    "Illinois was the primary driver. Opportunity remains in Missouri and Michigan."
)

doc1.add_heading("PRODUCT CATEGORY REVIEW", 1)
doc1.add_paragraph(
    "Technology: Strong performance, driven by Phones and Accessories sub-categories. "
    "Copiers delivered high-value orders despite lower unit volume. Margin held at 17%.\n\n"
    "Furniture: Revenue growth masked a profitability problem. The Tables sub-category "
    "ran at negative margin (-4.2%) for the full year due to excessive discounting and "
    "supplier cost increases. Chairs performed well. Bookcases stable.\n\n"
    "Office Supplies: Steady performance. Binders, Paper, and Storage drove volume. "
    "Labels and Fasteners showed declining demand -- possible product rationalisation "
    "candidate for 2018.\n\n"
    "[SECTION INCOMPLETE -- product manager inputs pending]"
)

doc1.add_heading("OUTLOOK 2018", 1)
doc1.add_paragraph(
    "Management targets for 2018:\n"
    "- Revenue: $2,500,000 (growth of ~8.8%)\n"
    "- Gross Margin: 14% (recovery from 12.5%)\n"
    "- Customer Retention: 72% (match industry benchmark)\n"
    "- Same-day delivery: Profitable unit economics by Q3 2018\n\n"
    "Key risks: Amazon Business competitive pressure, potential tariff increases on "
    "imported technology products, talent acquisition in analytics roles.\n\n"
    "[FULL 2018 PLAN -- SEE SEPARATE STRATEGY DOCUMENT]"
)

add_para(doc1, "\n\nDocument Classification: CONFIDENTIAL\nPrepared by: Finance & Strategy Team\nVersion: DRAFT 2.1 -- Not for external use", color=(150, 150, 150))
doc1.save(os.path.join(OUT, "annual_report_2017.docx"))
print("  Created: annual_report_2017.docx")


# DOCX 2 — Sales Team Meeting Notes (noisy, informal, partial data)
doc2 = Document()
doc2.add_heading("Sales Team Meeting Notes", 0)
add_para(doc2, "Date: November 14, 2O17  [typo in original -- OCR]", color=(200, 0, 0))
add_para(doc2, "Location: HQ Conference Room B (dial-in: 3 participants remote)")
add_para(doc2, "Facilitator: Sarah K.  |  Notes: Tom W. (may contain errors)")
doc2.add_paragraph()

doc2.add_heading("Attendees", 1)
doc2.add_paragraph(
    "Present: Sarah K. (VP Sales), Tom W. (Analytics), Raj P. (West), "
    "Linda H. (East), Dave M. (South -- remote, audio poor), "
    "Central rep: NO SHOW -- reschedule needed\n"
    "Apologies: CFO (conflict), Head of Ops (sick)"
)

doc2.add_heading("Agenda Items", 1)

doc2.add_heading("1. Q3 Performance Review", 2)
doc2.add_paragraph(
    "Sarah opened with Q3 numbers. West continues to lead -- Raj confirmed California "
    "pipeline strong going into Q4. Holiday season expected to boost Tech and Furniture.\n\n"
    "East performance: Linda flagged that New York Corporate accounts underperformed. "
    "3 large contracts delayed (total value ~$85,000) -- decision expected by Dec.\n\n"
    "South: Dave's audio was breaking up. Key points captured:\n"
    "  - [INAUDIBLE] ...discount approvals taking too long...\n"
    "  - Texas market share declining vs competitor\n"
    "  - Customer complaint rate up -- details unclear due to audio issues\n"
    "  - [MISSED -- could not hear, follow up with Dave directly]\n\n"
    "Central: No representative present. Review deferred to next meeting."
)

doc2.add_heading("2. Discount Policy Discussion", 2)
doc2.add_paragraph(
    "Tom presented analysis showing Furniture avg discount at 42% -- well above 20% target.\n"
    "Specific problem: Tables sub-category running at NEGATIVE margin.\n\n"
    "Discussion:\n"
    "  Raj: 'West customers expect high discounts on Furniture -- hard to reduce without losing deals'\n"
    "  Linda: 'East Corporate clients on 3-yr contracts -- can't change mid-term'\n"
    "  Sarah: 'We need to stop the bleeding. Finance is asking questions.'\n\n"
    "DECISION: New discount approval process from Dec 1:\n"
    "  - >20% discount requires manager sign-off\n"
    "  - >30% requires VP approval\n"
    "  - Tables: max 15% until margin recovers\n\n"
    "NOTE: This was agreed verbally -- policy document to follow. Tom to draft by Nov 21.\n"
    "[ACTION: Tom W. -- discount policy doc by 21 Nov 2017]"
)

doc2.add_heading("3. Q4 Targets", 2)
doc2.add_paragraph(
    "Targets set for Q4 (note: these were negotiated down from original board targets):\n"
    "  West   : $320,000   (original target was $350,000)\n"
    "  East   : $175,000\n"
    "  South  : $210,000   [Dave disputed -- says target too high given TX situation]\n"
    "  Central: $180,000\n"
    "  TOTAL  : $885,000   vs board target of $940,000\n\n"
    "Gap to board target: $55,000 -- to be recovered via stretch incentive programme\n"
    "[CONFIDENTIAL -- do not share targets outside this group]"
)

doc2.add_heading("4. New Product Lines (EMEA Expansion -- NOT Superstore)", 2)
doc2.add_paragraph(
    "Sarah briefly mentioned potential expansion to European market in 2019.\n"
    "This is SEPARATE from current US operations and NOT included in 2017/2018 plans.\n"
    "Preliminary discussions with UK distributor -- nothing confirmed.\n"
    "Product lines being considered for EU: different SKUs, different pricing, different "
    "supplier base. Not relevant to current reporting structure.\n"
    "[OUT OF SCOPE FOR CURRENT BI PROJECT -- flag if included in analysis]"
)

doc2.add_heading("5. Other Business", 2)
doc2.add_paragraph(
    "  - CRM upgrade scheduled Jan 2018 -- data migration risk flagged by Tom\n"
    "  - Holiday party: Dec 15, details TBC\n"
    "  - Next meeting: Dec 12 (note: conflicts with system maintenance window -- check IT)\n\n"
    "Meeting ended 15 mins early due to fire drill.\n\n"
    "TODO: Tom to circulate these notes for approval by Nov 17\n"
    "TODO: Central region rep to present at Dec meeting\n"
    "TODO: Dave to send South region corrected figures by email\n\n"
    "[NOTES NOT YET APPROVED -- TREAT AS DRAFT]"
)

add_para(doc2, "CONFIDENTIAL -- INTERNAL SALES USE ONLY", color=(180, 0, 0), bold=True)
doc2.save(os.path.join(OUT, "sales_team_meeting_notes.docx"))
print("  Created: sales_team_meeting_notes.docx")


# DOCX 3 — Supplier Contracts Summary (partial match, some garbage data)
doc3 = Document()
doc3.add_heading("Supplier Contracts Summary 2017-2018", 0)
add_para(doc3, "Procurement Department | STRICTLY CONFIDENTIAL", color=(180, 0, 0))
add_para(doc3, "Last Updated: [DATE PLACEHOLDER -- template not updated]", color=(150, 150, 150))
doc3.add_paragraph()

doc3.add_heading("Overview", 1)
doc3.add_paragraph(
    "This document summarises active supplier contracts relevant to TechRetail's "
    "product portfolio. Contracts cover Furniture, Technology, and Office Supplies "
    "categories sold across West, East, South, and Central US regions.\n\n"
    "Total active suppliers: 47\n"
    "Top 3 suppliers (by spend): 67% of total COGS\n"
    "Contracts due for renewal in 2018: 12\n\n"
    "IMPORTANT: Some contract values redacted pending legal review. "
    "Contact procurement for full figures."
)

doc3.add_heading("Key Suppliers by Category", 1)

doc3.add_heading("Furniture Suppliers", 2)
table2 = doc3.add_table(rows=1, cols=4)
table2.style = "Table Grid"
hdr2 = table2.rows[0].cells
for i, h in enumerate(["Supplier", "Contract Value", "Renewal Date", "Notes"]):
    hdr2[i].text = h
furniture_suppliers = [
    ("BushCorp Industries", "$1.2M/yr", "March 2018", "Main Bookcase & Desk supplier -- CRITICAL"),
    ("HonCo Seating", "$890K/yr", "June 2018", "Chairs sub-category -- good relationship"),
    ("GlobalFurn Ltd", "[REDACTED]", "Dec 2017 -- URGENT", "Tables supplier -- under review due to quality issues"),
    ("ErgoPro Solutions", "$340K/yr", "Sept 2018", "Standing desks -- growing demand"),
    ("Sauder Woodcraft", "$210K/yr", "Jan 2018", "Bookcases -- backup supplier"),
]
for r in furniture_suppliers:
    row = table2.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v

doc3.add_paragraph()
doc3.add_heading("Technology Suppliers", 2)
doc3.add_paragraph(
    "1. TechDistrib Inc. -- Primary tech distributor\n"
    "   Categories: Phones, Accessories, Machines, Copiers\n"
    "   Annual spend: ~$1.8M\n"
    "   Contract: Rolling 12-month, auto-renews Feb 2018\n"
    "   Note: Subject to tariff risk if China import duties increase\n\n"
    "2. Logitech Authorised Reseller Agreement\n"
    "   Categories: Accessories (keyboards, mice, webcams)\n"
    "   Annual spend: ~$420K\n"
    "   Contract: Expires Aug 2018 -- BEGIN RENEWAL TALKS Q1\n\n"
    "3. Samsung Business Direct\n"
    "   Categories: Phones, Accessories\n"
    "   Annual spend: [REDACTED]\n"
    "   Note: New agreement signed Nov 2017 -- details confidential\n\n"
    "4. XYZ Tech Imports (NON-US SUPPLIER -- flagged for compliance review)\n"
    "   Origin: Shenzhen, China\n"
    "   Products: Generic accessories, cables\n"
    "   Status: UNDER LEGAL REVIEW -- do not process new orders until cleared\n"
    "   [NOTE: This supplier NOT reflected in Superstore sales data]"
)

doc3.add_heading("Office Supplies Suppliers", 2)
doc3.add_paragraph(
    "Avery Dennison: Labels, binders -- $380K/yr -- stable\n"
    "3M Business: Post-its, tapes -- $290K/yr -- renews July 2018\n"
    "Staples Wholesale (yes, we buy from competitor's wholesale arm):\n"
    "   Paper products -- $520K/yr -- this is SENSITIVE -- do not disclose\n"
    "Dixon Ticonderoga: Art/craft supplies -- $45K/yr -- low priority\n"
    "GreenPaper Co: Recycled paper products -- $110K/yr -- new 2017 supplier\n"
    "   NOTE: Customer demand for eco products growing -- expand this relationship"
)

doc3.add_heading("Supplier Issues & Risks", 1)
doc3.add_paragraph(
    "CRITICAL: GlobalFurn (Tables supplier) -- contract expired Dec 2017, quality complaints\n"
    "HIGH: TechDistrib tariff exposure -- monitor US-China trade policy\n"
    "HIGH: XYZ Tech Imports compliance investigation -- all shipments on hold\n"
    "MEDIUM: Staples Wholesale relationship -- reputationally sensitive\n"
    "LOW: Several minor suppliers on 30-day rolling -- price volatility risk\n\n"
    "JUNK DATA BELOW -- IGNORE (test records not yet removed from system):\n"
    "Supplier X: $999,999,999 -- TEST\n"
    "Supplier NULL: N/A -- N/A -- DELETE\n"
    "AAAAA TEST AAAAA: $0.00 -- test entry 001 -- do not process\n"
    "[PROCUREMENT TO CLEAN TEST RECORDS BEFORE QUARTERLY AUDIT]"
)

add_para(doc3, "\nPREPARED BY: Procurement Dept  |  APPROVED BY: [PENDING]  |  VERSION: 1.4 DRAFT", color=(150, 150, 150))
doc3.save(os.path.join(OUT, "supplier_contracts_summary.docx"))
print("  Created: supplier_contracts_summary.docx")

print("  DOCX files done.")


# ─────────────────────────────────────────────
#  PDF FILES
# ─────────────────────────────────────────────

styles = getSampleStyleSheet()

def make_pdf(filename, title, sections):
    """sections: list of (heading_text, body_text) or (None, body_text) for plain para"""
    path = os.path.join(OUT, filename)
    doc = SimpleDocTemplate(path, pagesize=letter,
                            leftMargin=inch, rightMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    story = []

    title_style = ParagraphStyle("Title2", parent=styles["Title"], fontSize=18, spaceAfter=6)
    h1_style    = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=13, spaceBefore=12, spaceAfter=4)
    h2_style    = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=11, spaceBefore=8, spaceAfter=3)
    body_style  = ParagraphStyle("Body2", parent=styles["Normal"], fontSize=9, leading=13, spaceAfter=6)
    note_style  = ParagraphStyle("Note", parent=styles["Normal"], fontSize=8, leading=11,
                                 textColor=colors.orange, spaceAfter=4)

    story.append(Paragraph(title, title_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
    story.append(Spacer(1, 0.15*inch))

    for item in sections:
        if len(item) == 3 and item[2] == "note":
            story.append(Paragraph(item[1], note_style))
        elif item[0] == "H2":
            story.append(Paragraph(item[1], h2_style))
        elif item[0]:
            story.append(Paragraph(item[0], h1_style))
            if item[1]:
                story.append(Paragraph(item[1].replace("\n", "<br/>"), body_style))
        else:
            story.append(Paragraph(item[1].replace("\n", "<br/>"), body_style))

    doc.build(story)
    print(f"  Created: {filename}")


# PDF 1 — Executive Summary 2017
make_pdf("executive_summary_2017.pdf", "Executive Summary — FY2017 Performance", [
    ("Overview", "TechRetail Corporation closed FY2017 with total revenue of $2,297,200, "
     "representing 6.2% growth over FY2016. Gross margin declined slightly from 13.0% to 12.5%, "
     "driven primarily by Furniture discounting. The Technology category delivered the strongest "
     "margin performance at 17%, while Furniture lagged at 6.2%."),

    ("NOTE: These figures are management estimates pending final audit. "
     "Do not cite externally until audit completion (expected Feb 2018).", "", "note"),

    ("Regional Performance", "West Region: $725,457 | Best performing, led by California\n"
     "East Region: $678,781 | Steady growth, NY Corporate pipeline strong\n"
     "South Region: $501,239 | Below target; operational issues in Q4\n"
     "Central Region: $391,723 | Highest growth rate (12.2% YoY)\n\n"
     "[NOTE: South figures may include $12,000 adjustment not yet processed]"),

    ("H2", "Product Mix"),
    (None, "Technology: 36.4% of revenue — Phones, Accessories, Copiers led growth\n"
     "Furniture: 32.3% of revenue — Volume healthy but margin eroded\n"
     "Office Supplies: 31.3% of revenue — Stable, Binders and Paper dominant\n\n"
     "PROBLEM AREA: Tables sub-category generated NEGATIVE gross profit (-$17,725)\n"
     "Root cause: avg discount 38% vs 15% target + supplier cost increase mid-year"),

    ("H2", "Customer Segments"),
    (None, "Consumer: 51.5% of orders (48.3% of revenue) — largest volume segment\n"
     "Corporate: 30.0% of orders (32.5% of revenue) — highest avg order value\n"
     "Home Office: 18.5% of orders (19.2% of revenue) — fastest growing segment\n\n"
     "Customer retention rate: 68% [BELOW industry benchmark of 72%]\n"
     "NPS score: 34 (industry avg ~42) — improvement required"),

    ("H2", "2018 Targets (Preliminary)"),
    (None, "Revenue target: $2,500,000\nGross margin target: 14.0%\n"
     "Retention target: 72%\nKey initiatives: Furniture margin recovery, digital investment\n\n"
     "[FULL PLAN PENDING BOARD APPROVAL — THIS IS WORKING DRAFT ONLY]"),

    ("Key Risks",
     "1. Competitive: Amazon Business pricing pressure — HIGH\n"
     "2. Supplier: Tables supplier contract expired, quality issues — HIGH\n"
     "3. Technology: Import tariff uncertainty — MEDIUM\n"
     "4. Operational: Q4 South region outage recurrence risk — MEDIUM\n"
     "5. EU Expansion: Feasibility only — no commitments — LOW CURRENT RISK\n\n"
     "NOTE: This section references risks NOT in Superstore dataset "
     "(e.g. EU expansion, supplier compliance). Treat as context only."),

    (None, "\n\nDocument: CONFIDENTIAL DRAFT | Prepared: Finance & Strategy | "
     "Approved: PENDING | Version: 0.9"),
])


# PDF 2 — Operations Review Q3 2017 (more noise, mixed data)
make_pdf("operations_review_Q3_2017.pdf", "Operations Review — Q3 2017", [
    ("NOTE: This document was scanned and OCR-processed. "
     "Some figures may contain transcription errors (e.g. 0/O confusion, 1/l confusion). "
     "Cross-reference with source system before acting on specific numbers.", "", "note"),

    ("Order Fulfilment", "Total orders processed Q3 2O17: 2,573  [OCR: may be 2,573 or 2,518]\n"
     "On-time delivery rate: 91.3%\n"
     "Late deliveries: 223 orders\n"
     "Primary late causes:\n"
     "  - Carrier delays (FedEx): 142 cases\n"
     "  - Warehouse picking errors: 51 cases\n"
     "  - Address/data issues: 30 cases\n\n"
     "Region with highest delay rate: South (14.2% late)\n"
     "Region with lowest delay rate: West (5.1% late)"),

    ("H2", "Shipping Mode Analysis"),
    (None, "Standard Class : 5,832 orders  avg 5.0 days  cost/order $6.12\n"
     "Second Class   : 2,O14 orders  avg 3.1 days  cost/order $9.44   [O may be 0]\n"
     "First Class    : 1,512 orders  avg 2.0 days  cost/order $14.77\n"
     "Same Day       :   636 orders  avg 0.9 days  cost/order $23.18\n\n"
     "CONCERN: Same-day cost/order $23.18 vs revenue generated $18.40 — LOSS-MAKING\n"
     "Recommendation: Same-day restricted to orders >$150 only (currently no min)\n"
     "[ACTION: Ops to implement minimum order value for same-day — target date Oct 1]"),

    ("H2", "Returns & Refunds"),
    (None, "Total returns Q3: 334\nReturn rate: 3.3% of orders\n"
     "By category:\n"
     "  Furniture: 4.8% return rate — HIGHEST — mainly Tables and Bookcases\n"
     "  Technology: 2.1% return rate\n"
     "  Office Supplies: 1.9% return rate — LOWEST\n\n"
     "Top return reasons:\n"
     "  1. Damaged in transit (38%)\n"
     "  2. Wrong item shipped (24%)\n"
     "  3. Item not as described (19%)\n"
     "  4. Changed mind / no longer needed (14%)\n"
     "  5. Other / unknown (5%)  — [these need better categorisation in CRM]\n\n"
     "NOISE/GAP: Return data for Sept 25-30 missing — system migration in progress\n"
     "Estimated missing returns: ~40 orders — will be captured in Q4 report"),

    ("Warehouse Operations", "Warehouses: West (LA), East (NYC), South (Dallas), Central (Chicago)\n\n"
     "West LA:\n"
     "  Capacity utilisation: 78% — optimal\n"
     "  Staff: 34 FTE + 12 seasonal\n"
     "  Issues: None significant\n\n"
     "East NYC:\n"
     "  Capacity utilisation: 91% — AT RISK of overflow in Q4 peak\n"
     "  Staff: 28 FTE + 8 seasonal (8 more requested for Nov-Dec)\n"
     "  Issues: Rack storage system upgrade scheduled Oct — 3-day partial closure\n\n"
     "South Dallas:\n"
     "  Capacity utilisation: 65%\n"
     "  Staff: 22 FTE\n"
     "  Issues: 2 forklifts out of service — maintenance backlog\n"
     "  [NOTE: Dec system outage risk flagged — see separate IT report]\n\n"
     "Central Chicago:\n"
     "  Capacity utilisation: 70%\n"
     "  Staff: l9 FTE  [OCR: may be 19 or l9]\n"
     "  Issues: New WMS system implementation — some data gaps during transition"),

    ("NON-US OPERATIONS NOTE",
     "This report covers US operations ONLY.\n"
     "Canada pilot (3 stores, non-superstore product range) is tracked separately.\n"
     "Mexico distribution partnership: TERMINATED Q2 2017 — not in this data.\n"
     "Any references to international ops below are errors — please flag.\n\n"
     "[GARBAGE DATA FROM SYSTEM MIGRATION — IGNORE LINES BELOW]\n"
     "ZZZ_TEST_RECORD: warehouse=NULL location=UNKNOWN orders=9999999\n"
     "IMPORT_ERROR_ROW_4421: data corrupted during migration -- exclude from analysis\n"
     "---DELETE FROM HERE--- placeholder placeholder placeholder ---END DELETE---"),

    (None, "\n\nPrepared by: Operations Analytics Team\n"
     "Review cycle: Quarterly\nNext review: Q4 2017 (target date: Jan 15 2018)\n"
     "Distribution: Ops Leadership, Finance, CEO\n"
     "Version: FINAL 1.0  [Note: 'FINAL' but CFO comments still pending]"),
])


# PDF 3 — Board Presentation Notes (high-level, some unrelated strategic content)
make_pdf("board_presentation_notes_Dec2017.pdf", "Board Meeting Notes — December 2017", [
    ("NOTE: These are informal scribe notes, NOT official board minutes. "
     "Official minutes to be prepared by Company Secretary. "
     "Contents may be incomplete or inaccurate — verify before relying on.", "", "note"),

    ("Agenda Item 1: FY2017 Performance Update",
     "CFO presented headline numbers:\n"
     "  Revenue $2.3M — on track, slightly below $2.35M stretch target\n"
     "  Margin compressed — Furniture flagged as primary concern\n"
     "  West and East solid; South disappointing; Central improving\n\n"
     "Board questions:\n"
     "  Q: 'Why is Tables still loss-making after 2 years?' \n"
     "  A: CFO — supplier costs up, discount culture entrenched, fix in progress\n"
     "  Q: 'When does same-day become profitable?'\n"
     "  A: COO — minimum order threshold being implemented Q1 2018\n\n"
     "[Board member comment redacted per legal instruction]"),

    ("Agenda Item 2: 2018 Budget Approval",
     "Revenue target $2.5M APPROVED\n"
     "Capex budget $380,000 APPROVED (analytics platform $150K, warehouse upgrades $230K)\n"
     "Opex budget: DEFERRED — further detail required on marketing allocation\n\n"
     "Key conditions attached to approval:\n"
     "  1. Monthly margin reporting by category (not just overall)\n"
     "  2. Furniture turnaround plan submitted by Jan 31\n"
     "  3. Same-day delivery profitability report by Q2 board\n\n"
     "[CFO to circulate revised budget pack by Dec 20]"),

    ("H2", "Agenda Item 3: Strategic Options (CONFIDENTIAL)"),
    (None, "Three strategic options presented by external consultants [FIRM NAME REDACTED]:\n\n"
     "Option A: Organic growth — invest in digital, maintain current footprint\n"
     "Option B: Acquisition — acquire regional competitor in South or Central\n"
     "Option C: International expansion — EU market entry 2019\n\n"
     "Board discussion was lively. No decision taken.\n"
     "Option C (EU) was noted as high risk / high reward — feasibility study commissioned\n"
     "Option B: One target identified [NAME REDACTED] in South region — due diligence Q1\n\n"
     "IMPORTANT: Options B and C are NOT reflected in current Superstore dataset.\n"
     "Any revenue/product data from potential acquisition or EU operations would be\n"
     "in a DIFFERENT dataset and should not be mixed with current US Superstore figures."),

    ("Agenda Item 4: Risk Review",
     "Top risks discussed:\n"
     "  1. Amazon Business — board acknowledges structural threat\n"
     "     Mitigation: Focus on service quality + B2B relationships\n"
     "  2. Supplier concentration — board requested diversification plan\n"
     "  3. IT systems — aging ERP flagged as operational risk\n"
     "     ERP upgrade deferred to 2019 due to cost\n"
     "  4. Key person risk — CFO search ongoing (current CFO retiring Q2 2018)\n\n"
     "[ACTION: CEO to present risk mitigation plan at March board]"),

    ("AOB / Closing Notes",
     "Next board meeting: March 14, 2018\n"
     "AGM date: May 22, 2018\n"
     "Board dinner: Dec 20 — details by PA\n\n"
     "Meeting closed 18:47\n\n"
     "[SCRIBE NOTE: audio recording available but quality poor from 15:00 onwards.\n"
     "Agenda item 3 discussion may be incomplete in these notes.\n"
     "Board members should review draft minutes when circulated.]\n\n"
     "CONFIDENTIAL — BOARD USE ONLY — NOT FOR WIDER DISTRIBUTION"),
])

print("\nAll files generated successfully:")
for f in sorted(os.listdir(OUT)):
    size = os.path.getsize(os.path.join(OUT, f))
    print(f"  {f:50s}  {size/1024:6.1f} KB")
